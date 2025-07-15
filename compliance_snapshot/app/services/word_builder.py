from __future__ import annotations

from pathlib import Path
import sqlite3
import re
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .report_generator import (
    generate_hos_violations_summary,
    generate_hos_trend_analysis,
    generate_summary_insights,
    generate_trend_insights,
    generate_safety_inbox_summary,
    generate_safety_inbox_insights,
    generate_pc_usage_summary,
    generate_pc_usage_insights,
    generate_unassigned_driving_summary,
    generate_unassigned_driving_insights,
    generate_unassigned_segment_details,
    generate_speeding_analysis_summary,
    generate_speeding_analysis_insights,
    generate_missed_dvir_summary,
    generate_missed_dvir_insights,
    generate_dot_risk_assessment,
)

from .visualizations.chart_factory import (
    make_stacked_bar,
    make_trend_line,
    make_pc_usage_bar_chart,
    make_safety_events_bar,
    make_unassigned_segments_visual,
    make_speeding_pie_chart,
)


def load_data(wiz_id: str, table: str) -> pd.DataFrame:
    """Read ``table`` from the temporary SQLite DB for ``wiz_id``."""
    db_path = Path(f"/tmp/{wiz_id}/snapshot.db")
    con = sqlite3.connect(db_path)
    try:
        return pd.read_sql(f"SELECT * FROM {table}", con)
    finally:
        con.close()


def _strip_html(text: str) -> str:
    """Return ``text`` with simple HTML tags removed."""
    if not text:
        return ""
    text = re.sub(r"<span style=\"color:\s*red;\">([^<]+)</span>", r"\1", text)
    return re.sub(r"<[^>]+>", "", text)


def build_word(
    wiz_id: str,
    *,
    filters: dict | None = None,
    trend_end: str | None = None,
) -> Path:
    """Generate a Word version of the compliance snapshot report."""
    tmpdir = Path(f"/tmp/{wiz_id}")
    out_path = tmpdir / "ComplianceSnapshot.docx"

    df = load_data(wiz_id, "hos")
    if filters:
        for col, val in filters.items():
            if col in df.columns:
                df = df[df[col] == val]

    end_date = (
        pd.to_datetime(trend_end).date() if trend_end else pd.Timestamp.utcnow().date()
    )

    # Generate key statistics
    summary_data = generate_hos_violations_summary(df, end_date)
    trend_data = generate_hos_trend_analysis(df, end_date)

    summary_insights = _strip_html(generate_summary_insights(summary_data))
    trend_insights = _strip_html(generate_trend_insights(trend_data))

    # Load additional datasets for later sections
    def safe_load(name: str) -> pd.DataFrame:
        try:
            return load_data(wiz_id, name)
        except Exception:
            return pd.DataFrame()

    safety_df = safe_load("safety_inbox")
    pc_df = safe_load("personnel_conveyance")
    unassigned_df = safe_load("unassigned_hos")
    behaviors_df = safe_load("driver_behaviors")
    driver_safety_df = safe_load("driver_safety")
    mistdvi_df = safe_load("mistdvi")

    # Create charts
    bar_path = make_stacked_bar(df, tmpdir / "hos_bar.png")
    trend_path = make_trend_line(df, end_date, tmpdir / "hos_trend.png")
    safety_chart = make_safety_events_bar(safety_df, tmpdir / "safety_events.png")
    unassigned_chart = make_unassigned_segments_visual(
        unassigned_df, tmpdir / "unassigned_segments.png"
    )
    speeding_source = behaviors_df if not behaviors_df.empty else mistdvi_df
    speeding_chart = make_speeding_pie_chart(speeding_source, tmpdir / "speeding.png")
    pc_chart = make_pc_usage_bar_chart(pc_df, tmpdir / "pc_usage.png")

    # Build Word document
    doc = Document()
    doc.add_heading("DOT COMPLIANCE SNAPSHOT", 0)

    # Header section
    header_end = end_date
    start_date = header_end - pd.Timedelta(days=header_end.weekday())
    end_disp = start_date + pd.Timedelta(days=6)
    date_range_str = f"{start_date.strftime('%m/%d/%Y')} – {end_disp.strftime('%m/%d/%Y')}"

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).text = f"Date: {header_end.strftime('%B %d, %Y')}"
    table.cell(0, 1).text = "Location: All Regions"
    table.cell(1, 0).text = "Discussion Type: DOT Compliance"
    table.cell(1, 1).text = "Job Positions: Field Crew"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"DOT Fleet Safety Snapshot: {date_range_str}").bold = True

    # Dashboard section
    dashboard = doc.add_table(rows=3, cols=2)
    dashboard.style = "Table Grid"
    dashboard.rows[0].cells[0].paragraphs[0].add_run().add_picture(str(bar_path), width=Inches(3.2))
    dashboard.rows[0].cells[1].paragraphs[0].add_run().add_picture(str(trend_path), width=Inches(3.2))
    dashboard.rows[1].cells[0].paragraphs[0].add_run().add_picture(str(safety_chart), width=Inches(3.2))
    dashboard.rows[1].cells[1].paragraphs[0].add_run().add_picture(str(unassigned_chart), width=Inches(3.2))
    dashboard.rows[2].cells[0].paragraphs[0].add_run().add_picture(str(speeding_chart), width=Inches(3.2))
    dashboard.rows[2].cells[1].paragraphs[0].add_run().add_picture(str(pc_chart), width=Inches(3.2))

    doc.add_page_break()

    # HOS Violations Summary
    doc.add_heading("HOS Violations Summary", level=1)
    doc.add_paragraph(
        f"Total Violations: {summary_data['total_current']} ({summary_data['total_change']:+})",
        style="List Bullet",
    )
    for region, data in summary_data.get("by_region", {}).items():
        doc.add_paragraph(
            f"{region}: {data['current']} ({data['change']:+})",
            style="List Bullet 2",
        )
    doc.add_paragraph("Insights:", style="Intense Quote")
    doc.add_paragraph(summary_insights)

    doc.add_heading("4-Week Trend Analysis", level=1)
    doc.add_paragraph("Insights:", style="Intense Quote")
    doc.add_paragraph(trend_insights)

    doc.add_page_break()

    # Safety Inbox Events
    if not safety_df.empty:
        doc.add_heading("Safety Inbox Events", level=1)
        safety_data = generate_safety_inbox_summary(safety_df, end_date)
        doc.add_paragraph(
            f"Total Safety Events: {safety_data['total_current']} ({safety_data['total_change']:+})",
            style="List Bullet",
        )
        if safety_data.get("dismissed_count"):
            doc.add_paragraph(
                f"Dismissed: {safety_data['dismissed_count']}", style="List Bullet"
            )
        doc.add_paragraph("Insights:", style="Intense Quote")
        doc.add_paragraph(
            _strip_html(generate_safety_inbox_insights(safety_data))
        )
        doc.add_page_break()

    # Personal Conveyance usage
    if not pc_df.empty:
        doc.add_heading("Personal Conveyance (PC) Usage", level=1)
        pc_data = generate_pc_usage_summary(pc_df, end_date)
        doc.add_paragraph(
            f"Total PC Time: {pc_data['total_pc_time']} hours", style="List Bullet"
        )
        doc.add_paragraph("Top PC Users:")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        table.cell(0, 0).text = "Drivers"
        table.cell(0, 1).text = "Duration"
        for name, duration in pc_data["drivers_list"]:
            row = table.add_row().cells
            row[0].text = name
            row[1].text = duration
        row = table.add_row().cells
        row[0].text = "Grand Total"
        row[1].text = pc_data["grand_total"]
        doc.add_paragraph("Insights:", style="Intense Quote")
        doc.add_paragraph(_strip_html(generate_pc_usage_insights(pc_data)))
        doc.add_page_break()

    # Unassigned Driving segments
    if not unassigned_df.empty:
        unassigned_data = generate_unassigned_driving_summary(unassigned_df, end_date)
        doc.add_heading("Unassigned Driving", level=1)
        doc.add_paragraph("Insights:", style="Intense Quote")
        doc.add_paragraph(
            _strip_html(generate_unassigned_driving_insights(unassigned_data))
        )
        doc.add_paragraph("Unassigned Driving Segments", style="Intense Quote")
        doc.add_paragraph(
            _strip_html(generate_unassigned_segment_details(unassigned_data))
        )
        doc.add_page_break()

    # Driver Behavior & Speeding
    if not behaviors_df.empty or not driver_safety_df.empty:
        speeding_data = generate_speeding_analysis_summary(
            behaviors_df, driver_safety_df, end_date
        )
        doc.add_heading("Driver Behavior & Speeding", level=1)
        doc.add_paragraph("Insights:", style="Intense Quote")
        doc.add_paragraph(
            _strip_html(generate_speeding_analysis_insights(speeding_data))
        )
        doc.add_page_break()

    # Missed DVIRs
    if not mistdvi_df.empty:
        dvir_data = generate_missed_dvir_summary(mistdvi_df, end_date)
        doc.add_heading("Missed DVIRs (Pre/Post Trip Reports)", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Driver"
        hdr[1].text = "POST-TRIP"
        hdr[2].text = "PRE-TRIP"
        hdr[3].text = "Grand Total"
        for row_data in dvir_data["top_drivers"]:
            row = table.add_row().cells
            row[0].text = row_data["driver"]
            row[1].text = str(row_data["post_trip"])
            row[2].text = str(row_data["pre_trip"])
            row[3].text = str(row_data["total"])
        total_row = table.add_row().cells
        total_row[0].text = "Grand Total"
        total_row[1].text = str(dvir_data["total_post_trip"])
        total_row[2].text = str(dvir_data["total_pre_trip"])
        total_row[3].text = str(dvir_data["total_missed"])
        doc.add_paragraph("Insights:", style="Intense Quote")
        doc.add_paragraph(_strip_html(generate_missed_dvir_insights(dvir_data)))
        doc.add_page_break()

    # Overall DOT Risk Assessment
    risk = generate_dot_risk_assessment(
        summary_data,
        safety_df if not safety_df.empty else None,
        pc_df if not pc_df.empty else None,
        unassigned_df if not unassigned_df.empty else None,
        (
            generate_speeding_analysis_summary(behaviors_df, driver_safety_df, end_date)
            if (not behaviors_df.empty or not driver_safety_df.empty)
            else None
        ),
        dvir_data if "dvir_data" in locals() else None,
    )
    risk = _strip_html(risk.replace("####", "").replace("###", ""))
    doc.add_heading("Overall DOT Risk Assessment", level=1)
    doc.add_paragraph(risk)

    doc.save(out_path)
    return out_path
